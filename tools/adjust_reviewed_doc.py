from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path("/private/tmp/overall-accepted.docx")
OUTPUT = Path("/Users/wangyuan/Desktop/overallDesign/第三方通信管理软件总体设计方案_审阅调整版.docx")


def set_cell_paragraphs(cell, lines):
    """Replace cell prose while preserving table/cell geometry."""
    first = cell.paragraphs[0]
    first.text = lines[0]
    for para in list(cell.paragraphs[1:]):
        para._element.getparent().remove(para._element)
    anchor = first
    for text in lines[1:]:
        new_p = OxmlElement("w:p")
        anchor._p.addnext(new_p)
        anchor = type(first)(new_p, cell._element)
        anchor.text = text


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = type(paragraph)(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def remove_table_row(table, row_index):
    row = table.rows[row_index]
    table._tbl.remove(row._tr)


def replace_paragraph_text(paragraph, text, style=None, remove_numbering=False):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    if style:
        paragraph.style = style
    if remove_numbering:
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            ppr.remove(num_pr)
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


doc = Document(SOURCE)
paras = list(doc.paragraphs)

# 1) Confirmed design inputs: convert reviewer shorthand into normative wording.
set_cell_paragraphs(
    doc.tables[1].cell(2, 1),
    [
        "普通轮询协议点最快采集周期按 10 ms 设计。Modbus 主站和 OPC UA 客户端按配置周期主动轮询或订阅数据；"
        "Modbus 从站和 OPC UA 服务端被动响应外部请求。",
        "IEC 104 以变化上送为主，并保留周期上送、总召和死区等可配置参数；是否上送由点变化、死区阈值及链路状态共同决定。",
    ],
)
set_cell_paragraphs(
    doc.tables[1].cell(6, 1),
    [
        "点级数据路由按用途显式配置，不根据点名或 CardType 隐式推断。",
        "场景一（采集显示）：第三方数据映射为 AM/DM 点，写入共享内存并供上位画面读取。",
        "场景二（对外发送）：通信程序从共享内存读取 AM/DM 点，按目标协议编码后发送给第三方系统。",
        "场景三（进入控制器）：数据经 50 号虚拟站生成站间引用报文，由控制器接收并写入对应内部点。",
        "同一点需要同时用于画面显示和控制器时，应配置两条明确路由，并分别记录执行结果。",
    ],
)
set_cell_paragraphs(
    doc.tables[2].cell(5, 1),
    [
        "每个点保留 value、timestamp 和 status 三项运行属性。",
        "status 采用单一枚举：NotConnected（通道未建立）、Bad（通信已建立但点值无效）和 Good（通信正常且点值有效）。"
        "界面仅在通道已连接的前提下判定点质量，避免同时显示相互矛盾的通信状态和质量状态。",
    ],
)

# 2) Scope: ordinary points only; process alarms/SOE are not in phase one.
paras[22].text = (
    "与现有 StructManage 数据模型和 POSIX 共享内存兼容，形成统一的普通点读写、时间戳和状态语义。"
)
paras[28].text = (
    "一期只处理 AM/DM 普通点的数据采集、发布和状态维护，不接入现有报警服务和 SOE 服务；"
    "通信异常仅在本软件状态页与运行日志中呈现。报警与 SOE 接口作为后续扩展边界保留。"
)
paras[153].text = (
    "普通点更新按 PointKey 合并并保留最新值；写命令和审计事件不得静默丢弃。"
)
paras[154].text = (
    "报文与运行诊断采用异步文件写入：按通道和日期/大小轮转，达到磁盘水位后先停止原始报文落盘，"
    "但保留断链、恢复、连续超时和进程异常等关键事件记录，禁止磁盘慢写阻塞通信线程。"
)

# Align scope-related tables.
doc.tables[2].cell(2, 1).text = "连接现有 DCS 的普通点数据模型和共享内存；报警与 SOE 服务不在一期接入范围。"
doc.tables[3].cell(3, 1).text = (
    "Master/Client 与 Slave/Server；总召、周期/变化遥测遥信、自发上送和死区参数。"
)
doc.tables[3].cell(3, 2).text = "遥控、SOE、时钟同步、IEC 101、双机无扰切换和规约网关集群。"
doc.tables[4].cell(3, 1).text = "一期不写入 AlarmData/SOEData，仅保留适配接口位置供后续扩展。"
doc.tables[14].cell(5, 1).text = "一期不实现 SOE 映射；IEC104 带时标信息仅按普通点保存数值、源时间戳和状态。"

# 3) Move the user-inserted UI screenshot into the Qt UI chapter.
source_cell = doc.tables[6].cell(1, 0)
image_para = source_cell.paragraphs[0]
for run in image_para.runs:
    if not run._r.xpath(".//w:drawing"):
        run._r.getparent().remove(run._r)
image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
keep_with_next(image_para)

ui_intro = paras[159]
image_para._p.getparent().remove(image_para._p)
ui_intro._p.addnext(image_para._p)

caption = insert_after(image_para, "图 9-1  Qt5 Win98/DCS 风格工作台参考界面")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
keep_with_next(caption)
for run in caption.runs:
    run.bold = True
    run.font.size = Pt(9)

summary = insert_after(
    caption,
    "界面采用固定工作台结构：顶部负责工程和运行命令，左侧负责对象导航，中部承载主要编辑与监视页，"
    "右侧显示当前对象属性，底部集中呈现校验问题和运行日志。各区域通过 QDockWidget/QSplitter "
    "调整尺寸并保存布局，但关键运行状态始终可见。",
)

# Remove the two reviewer instruction lines left in the original architecture table.
for para in list(source_cell.paragraphs):
    remove_paragraph(para)
source_cell.add_paragraph(
    "界面参考图已移至第 9 章；本图仅保留总体架构，不混排界面审阅说明。"
)

# Replace the old “advantages” heading with the integrated functional partition.
paras[162].text = "9.1 工作台功能分区与交互基线"
paras[162].style = "Heading 2"

# Make the page matrix more concrete and consistent with the demo.
ui_rows = {
    1: ("系统总览", "通道状态、收发帧数、成功率、响应时间、点更新速率、资源使用和通信异常摘要。"),
    2: ("工程导航与组态", "左侧按工程→协议→通道→设备→扫描组→点表组织对象；双击对象在中部打开唯一页签。"),
    3: ("属性检查器", "右侧按基本、端点、时序、转换和诊断分组显示字段；使用类型化编辑器、单位、范围和即时校验。"),
    4: ("点表编辑", "中部虚拟化表格支持批量填充、复制粘贴、筛选、地址递增、导入导出、分包预览和错误定位。"),
    5: ("运行监视", "显示 value、timestamp、status、刷新频率及数据方向；一期状态统一为 NotConnected、Bad、Good。"),
    6: ("报文监视", "采用报文列表、十六进制原文和解析树三窗格；支持冻结、过滤、字节联动高亮和导出。"),
    7: ("问题与日志", "底部集中显示组态校验、连接错误、协议错误和运行日志；双击问题定位到对象或字段。"),
    8: ("版本与发布", "显示草稿差异、校验结果、发布版本和回滚入口；发布、写值和启停操作与普通查看明显区分。"),
}
for row_index, (name, detail) in ui_rows.items():
    doc.tables[43].cell(row_index, 0).text = name
    doc.tables[43].cell(row_index, 1).text = detail

# Packet export and runtime link-event recording.
doc.tables[41].cell(0, 0).text = (
    "DISABLED → STOPPED → STARTING → CONNECTING/HANDSHAKING → RUNNING\n"
    "                         ↑              │                    │\n"
    "                         └─ RETRY_WAIT ←┴─ DEGRADED/FAILED ←─┘\n\n"
    "运行时持续检测连接、握手、心跳、收发超时和进程心跳。每次状态迁移均生成事件并异步写入诊断文件，"
    "至少包含通道 ID、协议、前后状态、原因码、对端、时间戳、持续时间、重试次数和配置版本。"
)
doc.tables[42].cell(2, 1).text = (
    "十六进制 + ASCII；收发不同颜色；显示帧边界；支持复制单帧，并可按当前筛选结果批量导出。"
)
doc.tables[42].cell(4, 1).text = (
    "按通道、方向、协议、地址/功能码/TI、成功失败、时间区间和文本过滤；导出操作必须复用当前过滤条件。"
)
doc.tables[42].cell(6, 1).text = (
    "按权限启动/停止；支持 pcapng 和应用层 CSV/JSONL 自定义记录，导出内容包含时间、方向、通道、对端、"
    "原始字节、解析摘要和结果，供 Wireshark、脚本或离线分析工具使用；文件按大小/时间轮转并自动脱敏凭据。"
)
doc.tables[42].cell(7, 1).text = (
    "一键导出配置摘要、版本、通道状态、断链/恢复事件、近端日志、统计与有限报文，不导出密钥；"
    "诊断包携带时间范围和校验清单，便于问题复现与归档。"
)

# Rename subsequent subsections after removal/integration of the old 9.1 wording.
paras[164].text = "9.2 实现约束"
paras[171].text = "9.3 协议界面插件与工厂"

# 4) Simplified point state model.
paras[181].text = "11. 时间戳与普通点状态"
doc.tables[49].cell(2, 1).text = (
    "每点保留一个 timestamp。协议提供源时标时使用源时标，否则使用本机接收时刻；内部统一存储 UTC。"
)
doc.tables[49].cell(3, 1).text = (
    "统一为 NotConnected、Bad、Good。NotConnected 表示通道未建立；Bad 表示通道已建立但本点数据无效；"
    "Good 表示通信正常且点值有效。"
)
doc.tables[49].cell(4, 0).text = "状态更新"
doc.tables[49].cell(4, 1).text = (
    "断线时通道内全部点置为 NotConnected；解析失败或超时但通道仍在线时相关点置为 Bad；"
    "收到首个有效值后恢复为 Good。"
)
remove_table_row(doc.tables[49], 6)
remove_table_row(doc.tables[49], 5)

# 5) No routine sudo. Keep operational alerts distinct from DCS process alarms.
paras[185].text = (
    "生产运行、通道启停、组态发布和日志采集均不得依赖交互式 sudo。安装阶段由部署脚本一次性创建专用账户、"
    "设备组、目录和 systemd 权限；运行期仅使用预授权能力，确需提权的维护动作必须独立封装并记录审计。"
)
paras[186].text = (
    "最小权限运行：服务、GUI 和 Python worker 使用受控账户/权限；串口通过专用设备组授权，证书和配置目录"
    "仅授予必要读写权限，禁止为方便部署授予全局 sudo。"
)
paras[187].text = (
    "RBAC 细分查看、组态、发布、启停、写值、脚本、证书和审计导出权限。"
)

# 6) Four developers and five repositories, using explicit integration boundaries.
replace_paragraph_text(paras[221], "实施计划与交付物", style="Heading 1")
plan_text = (
    "项目计划由 4 名开发人员并行实施，代码按 5 个 Git 仓库管理：1 个主界面/集成仓库和 4 个通信模块仓库"
    "（Modbus、IEC104、OPC UA、Python/UDP）。主仓库只通过稳定接口和固定版本引用通信模块，建议使用 Git "
    "submodule 或包版本清单锁定 commit，禁止直接复制源码形成不可追踪的嵌套仓库。人员可分别负责界面与集成、"
    "Modbus、IEC104、OPC UA/Python；公共 DTO、错误码和测试向量由主仓库统一发布，确保各模块互不干扰且可独立构建测试。"
)
replace_paragraph_text(paras[222], plan_text, style="Normal", remove_numbering=True)
doc.tables[58].cell(4, 2).text = "链路/APCI/ASDU、总召、周期与变化上送、死区参数、普通点映射和仿真测试"
doc.tables[50].cell(1, 1).text = (
    "运行时检测 socket 断开、心跳失败和连续超时；立即记录断链事件文件并将通道置为重连状态。"
    "采用指数退避 + 抖动并设置最大退避；恢复后按协议重新握手/总召/订阅，同时记录恢复时间、"
    "中断持续时长和重试次数。"
)
doc.tables[50].cell(4, 1).text = (
    "达到磁盘高水位时停止原始报文落盘，保留通信和关键断链/恢复事件；采用文件轮转、保留期限及应急预留空间。"
)

# Remove remaining phase-one contradictions in test and adapter descriptions.
doc.tables[53].cell(2, 2).text = (
    "Iec104Channel、点映射、总召策略、周期/变化上送、死区参数、质量转换、统计和审计；"
    "一期不实现遥控与 SOE 映射"
)
doc.tables[57].cell(3, 1).text = (
    "与模拟器和真实设备互联；断线恢复；组态热更新；写后读；IEC104 周期/变化上送及源时间戳。"
)
paras[239].text = (
    "IEC104 周期/变化上送、死区和总召的现场语义；OPC UA 证书签发和信任流程。"
)

# Preserve original formatting and save as a separate adjusted deliverable.
for element in list(doc.element.xpath(".//*[local-name()='pPrChange' or local-name()='rPrChange' or "
                                      "local-name()='tblPrChange' or local-name()='trPrChange' or "
                                      "local-name()='tcPrChange' or local-name()='sectPrChange']")):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
track_revisions = doc.settings.element.find(qn("w:trackRevisions"))
if track_revisions is not None:
    doc.settings.element.remove(track_revisions)
doc.save(OUTPUT)
print(OUTPUT)
